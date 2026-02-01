"""
Document Tools
Tools for generating legal documents.
"""

from typing import Dict, List, Any, Optional
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.firestore_service import get_firestore_service
from services.storage_service import get_storage_service


async def generate_legal_memo(
    session_id: str,
    contract_id: Optional[str],
    title: str,
    subject: str,
    analysis: str,
    findings: List[Dict],
    recommendations: List[str],
    prepared_by: str = "LegalMind AI",
) -> Dict[str, Any]:
    """Generate a formal legal memorandum document.
    
    Args:
        session_id: Associated session ID
        contract_id: Associated contract ID (optional)
        title: Memo title
        subject: Subject line
        analysis: Main analysis text
        findings: List of findings with title and description
        recommendations: List of recommendations
        prepared_by: Author name
        
    Returns:
        Generated document info with download URL
    """
    firestore = get_firestore_service()
    storage = get_storage_service()
    
    # Create Word document
    doc = Document()
    
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "CONFIDENTIAL LEGAL MEMORANDUM"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Memo header info
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("TO:", "Legal Review"),
        ("FROM:", prepared_by),
        ("DATE:", datetime.now().strftime("%B %d, %Y")),
        ("RE:", subject),
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading("EXECUTIVE SUMMARY", level=1)
    doc.add_paragraph(analysis[:500] + "..." if len(analysis) > 500 else analysis)
    
    # Findings
    doc.add_heading("KEY FINDINGS", level=1)
    for i, finding in enumerate(findings, 1):
        doc.add_heading(f"{i}. {finding.get('title', 'Finding')}", level=2)
        doc.add_paragraph(finding.get('description', ''))
        
        if finding.get('severity'):
            severity_para = doc.add_paragraph()
            severity_para.add_run("Severity: ").bold = True
            severity_para.add_run(finding['severity'].upper())
    
    # Full Analysis
    doc.add_heading("DETAILED ANALYSIS", level=1)
    doc.add_paragraph(analysis)
    
    # Recommendations
    doc.add_heading("RECOMMENDATIONS", level=1)
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f"{i}. {rec}", style='List Number')
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run("This memorandum was prepared by LegalMind AI and is for informational purposes only. ")
    footer_para.add_run("It does not constitute legal advice. Please consult with a licensed attorney.")
    footer_para.italic = True
    
    # Save to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    # Upload to Cloud Storage
    doc_id = f"memo_{session_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    file_url = await storage.upload_generated_document(
        doc_bytes,
        doc_id,
        "memo",
        ".docx"
    )
    
    # Save document record
    await firestore.create_generated_document(
        session_id=session_id,
        contract_id=contract_id,
        document_type="legal_memo",
        title=title,
        file_url=file_url,
        content_summary=subject,
    )
    
    # Generate signed URL for download
    signed_url = await storage.get_signed_url(file_url, expiration_minutes=60)
    
    return {
        "status": "success",
        "document_type": "legal_memo",
        "title": title,
        "file_url": file_url,
        "download_url": signed_url,
        "expires_in_minutes": 60,
    }


async def generate_contract_summary(
    session_id: str,
    contract_id: str,
) -> Dict[str, Any]:
    """Generate an executive summary of a contract.
    
    Args:
        session_id: Associated session ID
        contract_id: The contract ID
        
    Returns:
        Generated summary document info
    """
    firestore = get_firestore_service()
    storage = get_storage_service()
    
    # Get contract data
    contract = await firestore.get_contract(contract_id)
    if not contract:
        return {
            "status": "error",
            "message": f"Contract {contract_id} not found"
        }
    
    # Get clauses
    clauses = await firestore.get_clauses_for_contract(contract_id)
    
    # Create document
    doc = Document()
    
    # Title
    doc.add_heading("CONTRACT SUMMARY", level=0)
    
    # Contract Info
    doc.add_heading("Contract Information", level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("Title:", contract.get("title", "N/A")),
        ("Type:", contract.get("contract_type", "N/A")),
        ("Status:", contract.get("status", "N/A")),
        ("Risk Level:", contract.get("risk_level", "N/A")),
        ("Compliance:", contract.get("compliance_status", "N/A")),
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)
    
    # Parties
    parties = contract.get("parties", [])
    if parties:
        doc.add_heading("Parties", level=1)
        for party in parties:
            doc.add_paragraph(f"• {party.get('name', 'Unknown')} ({party.get('role', 'Party')})")
    
    # Key Dates
    key_dates = contract.get("key_dates", [])
    if key_dates:
        doc.add_heading("Key Dates", level=1)
        for date_info in key_dates:
            doc.add_paragraph(f"• {date_info.get('date', 'N/A')}: {date_info.get('description', '')}")
    
    # Risk Summary
    risk_score = contract.get("overall_risk_score")
    if risk_score is not None:
        doc.add_heading("Risk Assessment", level=1)
        doc.add_paragraph(f"Overall Risk Score: {risk_score}/100 ({contract.get('risk_level', 'Unknown')})")
        
        risk_findings = contract.get("risk_findings", [])
        if risk_findings:
            doc.add_paragraph("Key Risk Findings:")
            for finding in risk_findings[:5]:  # Top 5 findings
                doc.add_paragraph(f"• {finding.get('risk_type', 'Unknown')}: {finding.get('description', '')}")
    
    # Clause Summary
    if clauses:
        doc.add_heading("Clause Summary", level=1)
        
        # Group by type
        clause_types = {}
        for clause in clauses:
            ctype = clause.get("clause_type", "general")
            if ctype not in clause_types:
                clause_types[ctype] = []
            clause_types[ctype].append(clause)
        
        for ctype, type_clauses in clause_types.items():
            doc.add_heading(ctype.replace("_", " ").title(), level=2)
            high_risk = [c for c in type_clauses if c.get("risk_level") in ["high", "critical"]]
            doc.add_paragraph(f"Count: {len(type_clauses)} | High Risk: {len(high_risk)}")
    
    # Save and upload
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    doc_id = f"summary_{contract_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    file_url = await storage.upload_generated_document(
        doc_bytes,
        doc_id,
        "summary",
        ".docx"
    )
    
    await firestore.create_generated_document(
        session_id=session_id,
        contract_id=contract_id,
        document_type="contract_summary",
        title=f"Summary: {contract.get('title', 'Contract')}",
        file_url=file_url,
    )
    
    signed_url = await storage.get_signed_url(file_url, expiration_minutes=60)
    
    return {
        "status": "success",
        "document_type": "contract_summary",
        "contract_id": contract_id,
        "file_url": file_url,
        "download_url": signed_url,
        "expires_in_minutes": 60,
    }


async def generate_risk_report(
    session_id: str,
    contract_id: str,
) -> Dict[str, Any]:
    """Generate a detailed risk assessment report.
    
    Args:
        session_id: Associated session ID
        contract_id: The contract ID
        
    Returns:
        Generated risk report info
    """
    firestore = get_firestore_service()
    storage = get_storage_service()
    
    contract = await firestore.get_contract(contract_id)
    if not contract:
        return {
            "status": "error",
            "message": f"Contract {contract_id} not found"
        }
    
    clauses = await firestore.get_clauses_for_contract(contract_id)
    
    doc = Document()
    
    # Title
    doc.add_heading("RISK ASSESSMENT REPORT", level=0)
    doc.add_paragraph(f"Contract: {contract.get('title', 'Unknown')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    
    # Overall Risk Score
    doc.add_heading("Risk Score Summary", level=1)
    risk_score = contract.get("overall_risk_score", 0)
    risk_level = contract.get("risk_level", "unknown")
    
    # Add visual indicator
    score_para = doc.add_paragraph()
    score_para.add_run(f"Overall Risk Score: {risk_score}/100").bold = True
    score_para.add_run(f" ({risk_level.upper()})")
    
    # Risk distribution
    doc.add_heading("Clause Risk Distribution", level=1)
    
    risk_counts = {"low": 0, "medium": 0, "high": 0, "critical": 0}
    for clause in clauses:
        level = clause.get("risk_level", "low")
        if level in risk_counts:
            risk_counts[level] += 1
    
    dist_table = doc.add_table(rows=5, cols=2)
    dist_table.style = 'Table Grid'
    
    dist_data = [
        ("Risk Level", "Clause Count"),
        ("Low", str(risk_counts["low"])),
        ("Medium", str(risk_counts["medium"])),
        ("High", str(risk_counts["high"])),
        ("Critical", str(risk_counts["critical"])),
    ]
    
    for i, (label, value) in enumerate(dist_data):
        row = dist_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    # Risk Findings
    findings = contract.get("risk_findings", [])
    if findings:
        doc.add_heading("Risk Findings", level=1)
        
        for finding in findings:
            doc.add_heading(finding.get("risk_type", "Unknown").replace("_", " ").title(), level=2)
            doc.add_paragraph(finding.get("description", ""))
            doc.add_paragraph(f"Score Impact: {finding.get('score', 0)} | Severity: {finding.get('severity', 'unknown').upper()}")
            
            matches = finding.get("matches", [])
            if matches:
                doc.add_paragraph("Matching Text:")
                for match in matches[:3]:
                    doc.add_paragraph(f"  • \"{match.get('pattern', '')}\" - {match.get('context', '')[:200]}", style='List Bullet')
    
    # High Risk Clauses
    high_risk_clauses = [c for c in clauses if c.get("risk_level") in ["high", "critical"]]
    if high_risk_clauses:
        doc.add_heading("High Risk Clauses", level=1)
        
        for clause in high_risk_clauses:
            doc.add_heading(f"Section {clause.get('section_number', 'N/A')}: {clause.get('clause_type', 'Unknown').replace('_', ' ').title()}", level=2)
            doc.add_paragraph(f"Risk Level: {clause.get('risk_level', 'Unknown').upper()}")
            doc.add_paragraph(f"Risk Explanation: {clause.get('risk_explanation', 'N/A')}")
            
            if clause.get("content"):
                doc.add_paragraph("Clause Text:")
                doc.add_paragraph(clause["content"][:500] + "..." if len(clause.get("content", "")) > 500 else clause["content"])
    
    # Recommendations
    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph("Based on the risk assessment, the following actions are recommended:")
    
    # Generate recommendations from findings
    rec_num = 1
    for finding in findings:
        if finding.get("severity") in ["high", "critical"]:
            doc.add_paragraph(f"{rec_num}. Address {finding.get('risk_type', 'risk').replace('_', ' ')} issues", style='List Number')
            rec_num += 1
    
    # Save and upload
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    doc_id = f"risk_report_{contract_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    file_url = await storage.upload_generated_document(
        doc_bytes,
        doc_id,
        "risk_report",
        ".docx"
    )
    
    await firestore.create_generated_document(
        session_id=session_id,
        contract_id=contract_id,
        document_type="risk_report",
        title=f"Risk Report: {contract.get('title', 'Contract')}",
        file_url=file_url,
    )
    
    signed_url = await storage.get_signed_url(file_url, expiration_minutes=60)
    
    return {
        "status": "success",
        "document_type": "risk_report",
        "contract_id": contract_id,
        "file_url": file_url,
        "download_url": signed_url,
        "expires_in_minutes": 60,
    }


async def list_generated_documents(
    session_id: Optional[str] = None,
    contract_id: Optional[str] = None,
    document_type: Optional[str] = None,
    limit: int = 50,
) -> Dict[str, Any]:
    """List generated documents with filters.
    
    Args:
        session_id: Filter by session
        contract_id: Filter by contract
        document_type: Filter by type
        limit: Maximum results
        
    Returns:
        List of documents
    """
    firestore = get_firestore_service()
    storage = get_storage_service()
    
    documents = await firestore.list_documents(
        session_id=session_id,
        contract_id=contract_id,
        document_type=document_type,
        limit=limit,
    )
    
    # Add download URLs
    for doc in documents:
        if doc.get("file_url"):
            doc["download_url"] = await storage.get_signed_url(
                doc["file_url"],
                expiration_minutes=60
            )
    
    return {
        "status": "success",
        "documents": documents,
        "count": len(documents),
    }


# Tool definitions for Gemini function calling
DOCUMENT_TOOLS = [
    {
        "name": "generate_legal_memo",
        "description": "Generate a formal legal memorandum document with findings and recommendations.",
        "parameters": {
            "type": "object",
            "properties": {
                "session_id": {
                    "type": "string",
                    "description": "Current session ID"
                },
                "contract_id": {
                    "type": "string",
                    "description": "Associated contract ID (optional)"
                },
                "title": {
                    "type": "string",
                    "description": "Memo title"
                },
                "subject": {
                    "type": "string",
                    "description": "Memo subject line"
                },
                "analysis": {
                    "type": "string",
                    "description": "Main analysis text"
                },
                "findings": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "description": {"type": "string"},
                            "severity": {"type": "string"}
                        }
                    },
                    "description": "List of findings"
                },
                "recommendations": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of recommendations"
                },
                "prepared_by": {
                    "type": "string",
                    "description": "Author name (default: LegalMind AI)"
                }
            },
            "required": ["session_id", "title", "subject", "analysis", "findings", "recommendations"]
        },
        "handler": generate_legal_memo
    },
    {
        "name": "generate_contract_summary",
        "description": "Generate an executive summary document for a contract.",
        "parameters": {
            "type": "object",
            "properties": {
                "session_id": {
                    "type": "string",
                    "description": "Current session ID"
                },
                "contract_id": {
                    "type": "string",
                    "description": "The contract ID to summarize"
                }
            },
            "required": ["session_id", "contract_id"]
        },
        "handler": generate_contract_summary
    },
    {
        "name": "generate_risk_report",
        "description": "Generate a detailed risk assessment report for a contract.",
        "parameters": {
            "type": "object",
            "properties": {
                "session_id": {
                    "type": "string",
                    "description": "Current session ID"
                },
                "contract_id": {
                    "type": "string",
                    "description": "The contract ID"
                }
            },
            "required": ["session_id", "contract_id"]
        },
        "handler": generate_risk_report
    },
    {
        "name": "list_generated_documents",
        "description": "List previously generated documents with optional filters.",
        "parameters": {
            "type": "object",
            "properties": {
                "session_id": {
                    "type": "string",
                    "description": "Filter by session ID"
                },
                "contract_id": {
                    "type": "string",
                    "description": "Filter by contract ID"
                },
                "document_type": {
                    "type": "string",
                    "description": "Filter by document type",
                    "enum": ["legal_memo", "contract_summary", "risk_report", "compliance_report"]
                },
                "limit": {
                    "type": "integer",
                    "description": "Maximum results (default 50)"
                }
            }
        },
        "handler": list_generated_documents
    }
]


def get_document_tools() -> List[Dict[str, Any]]:
    """Get all document tool definitions."""
    return DOCUMENT_TOOLS


# Export for tool registry
TOOL_DEFINITIONS = DOCUMENT_TOOLS
